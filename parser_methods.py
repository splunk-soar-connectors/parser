# File: parser_methods.py
#
# Copyright (c) 2017-2025 Splunk Inc.
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software distributed under
# the License is distributed on an "AS IS" BASIS, WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND,
# either express or implied. See the License for the specific language governing permissions
# and limitations under the License.
import csv
import re
import struct
import threading
import time
import zipfile
from html import unescape
from io import StringIO
from typing import TYPE_CHECKING, Any, Optional, TypedDict, Union, cast
from urllib.parse import urlparse

import docx
import phantom.app as phantom
import phantom.utils as ph_utils
from bs4 import BeautifulSoup
from bs4.dammit import UnicodeDammit
from django.core.validators import URLValidator
from docx.opc.constants import RELATIONSHIP_TYPE as REL_TYPE
from docx.opc.part import Part as DocxPart
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfdocument import PDFDocument, PDFEncryptionError, PDFPasswordIncorrect
from pdfminer.pdfinterp import PDFPageInterpreter, PDFResourceManager
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfparser import PDFParser
from pdfminer.pdftypes import PDFObjectNotFound, PDFObjRef, PDFStream
from pdfminer.psparser import PSKeyword, PSLiteral
from pdfminer.utils import isnumber


if TYPE_CHECKING:
    from phantom.action_result import ActionResult
    from phantom.base_connector import BaseConnector

_container_common = {"run_automation": False}  # Don't run any playbooks, when this artifact is added


URI_REGEX = r"h(?:tt|xx)p[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+#]|[!*\(\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+"
EMAIL_REGEX = r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b"
EMAIL_REGEX2 = r'".*"@[A-Z0-9.-]+\.[A-Z]{2,}\b'
HASH_REGEX = r"\b[0-9a-fA-F]{32}\b|\b[0-9a-fA-F]{40}\b|\b[0-9a-fA-F]{64}\b"
IP_REGEX = r"\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}"
IPV6_REGEX = (
    r"\s*((([0-9A-Fa-f]{1,4}:){7}([0-9A-Fa-f]{1,4}|:))|"
    r"(([0-9A-Fa-f]{1,4}:){6}(:[0-9A-Fa-f]{1,4}|((25[0-5]|2[0-4]\d|1\d\d|[1-9]?\d)"
    r"(\.(25[0-5]|2[0-4]\d|1\d\d|[1-9]?\d)){3})|:))"
    r"|(([0-9A-Fa-f]{1,4}:){5}(((:[0-9A-Fa-f]{1,4}){1,2})"
    r"|:((25[0-5]|2[0-4]\d|1\d\d|[1-9]?\d)(\.(25[0-5]|2[0-4]\d|1\d\d|[1-9]?\d)){3})|:))|"
    r"(([0-9A-Fa-f]{1,4}:){4}(((:[0-9A-Fa-f]{1,4}){1,3})"
    r"|((:[0-9A-Fa-f]{1,4})?:((25[0-5]|2[0-4]\d|1\d\d|[1-9]?\d)"
    r"(\.(25[0-5]|2[0-4]\d|1\d\d|[1-9]?\d)){3}))|:))|"
    r"(([0-9A-Fa-f]{1,4}:){3}(((:[0-9A-Fa-f]{1,4}){1,4})"
    r"|((:[0-9A-Fa-f]{1,4}){0,2}:((25[0-5]|2[0-4]\d|1\d\d|[1-9]?\d)"
    r"(\.(25[0-5]|2[0-4]\d|1\d\d|[1-9]?\d)){3}))|:))|"
    r"(([0-9A-Fa-f]{1,4}:){2}(((:[0-9A-Fa-f]{1,4}){1,5})"
    r"|((:[0-9A-Fa-f]{1,4}){0,3}:((25[0-5]|2[0-4]\d|1\d\d|[1-9]?\d)"
    r"(\.(25[0-5]|2[0-4]\d|1\d\d|[1-9]?\d)){3}))|:))|"
    r"(([0-9A-Fa-f]{1,4}:){1}(((:[0-9A-Fa-f]{1,4}){1,6})"
    r"|((:[0-9A-Fa-f]{1,4}){0,4}:((25[0-5]|2[0-4]\d|1\d\d|[1-9]?\d)"
    r"(\.(25[0-5]|2[0-4]\d|1\d\d|[1-9]?\d)){3}))|:))|"
    r"(:(((:[0-9A-Fa-f]{1,4}){1,7})|((:[0-9A-Fa-f]{1,4}){0,5}:((25[0-5]|2[0-4]\d|1\d\d|[1-9]?\d)"
    r"(\.(25[0-5]|2[0-4]\d|1\d\d|[1-9]?\d)){3}))|:)))(%.+)?\s*"
)
DOMAIN_REGEX = r"(?!:\/\/)((?:[a-zA-Z0-9-_]+\.)*[a-zA-Z0-9][a-zA-Z0-9-_]+\.[a-zA-Z]{2,11})"

ESCAPE = set(map(ord, '&<>"'))


class Artifact(TypedDict):
    source_data_identifier: int
    cef: dict[str, Any]
    name: str


class FileInfo(TypedDict):
    type: str
    path: str
    name: str
    id: Optional[str]


def _extract_domain_from_url(url: str) -> Optional[str]:
    domain = phantom.get_host_from_url(url)
    if domain and not _is_ip(domain):
        return domain
    return None


def _is_ip(input_ip: str) -> bool:
    if ph_utils.is_ip(input_ip):
        return True

    if is_ipv6(input_ip):
        return True

    return False


def _is_url(input_url: str) -> bool:
    validate_url = URLValidator(schemes=["http", "https"])
    try:
        validate_url(input_url)
        return True
    except Exception:
        return False


def is_ipv6(input_ip: str) -> bool:
    return bool(re.match(IPV6_REGEX, input_ip))


def _refang_url(url: str) -> str:
    parsed = urlparse(url)
    scheme = parsed.scheme

    # Replace hxxp/hxxps with http/https
    if scheme == "hxxp":
        parsed = parsed._replace(scheme="http")
    elif scheme == "hxxps":
        parsed = parsed._replace(scheme="https")

    refang_url = parsed.geturl()
    return refang_url


def _clean_url(url: str) -> str:
    url = url.strip(">),.]\r\n")

    # Check before splicing, find returns -1 if not found
    # _and_ you will end up splicing on -1 (incorrectly)
    if "<" in url:
        url = url[: url.find("<")]

    if ">" in url:
        url = url[: url.find(">")]

    url = _refang_url(url)
    return url


def _get_error_message_from_exception(e: Exception) -> tuple[Union[str, int], str]:
    """This method is used to get appropriate error message from the exception.
    :param e: Exception object
    :return: error message
    """
    error_msg = "Unknown error occured. Please check asset configuration and/or action parameters"
    error_code = "Error code unavailable"
    try:
        if hasattr(e, "args"):
            if len(e.args) > 1:
                error_code = e.args[0]
                error_msg = e.args[1]
            elif len(e.args) == 1:
                error_code = "Error code unavailable"
                error_msg = e.args[0]
            else:
                error_msg = "Unknown error occured. Please check asset configuration and/or action parameters"
                error_code = "Error code unavailable"
        else:
            error_code = "Error code unavailable"
            error_msg = "Error message unavailable. Please check the action parameters."
    except Exception:
        error_code = "Error code unavailable"
        error_msg = "Error message unavailable. Please check the action parameters."

    return error_code, error_msg


class TextIOCParser:
    BASE_PATTERNS = [
        {
            "cef": "sourceAddress",  # Name of CEF field
            "pattern": IP_REGEX,  # Regex to match
            "name": "IP Artifact",  # Name of artifact
            "validator": _is_ip,  # Additional function to verify matched string (Should return true or false)
        },
        {
            "cef": "sourceAddress",
            "pattern": IPV6_REGEX,
            "name": "IP Artifact",
            "validator": _is_ip,
        },
        {
            "cef": "requestURL",
            "pattern": URI_REGEX,
            "name": "URL Artifact",
            "clean": _clean_url,  # Additional cleaning of data from regex (Should return a string)
            "validator": _is_url,
        },
        {"cef": "fileHash", "pattern": HASH_REGEX, "name": "Hash Artifact"},
        {"cef": "email", "pattern": EMAIL_REGEX, "name": "Email Artifact"},
        {"cef": "email", "pattern": EMAIL_REGEX2, "name": "Email Artifact"},
    ]
    DOMAIN_PATTERN = {
        "cef": "destinationDnsDomain",
        "pattern": DOMAIN_REGEX,
        "name": "Domain Artifact",
    }  # Name of CEF field  # Regex to match

    URL_DOMAIN_SUBTYPES_DICT = {
        "subtypes": [  # Additional IOCs to find in a matched one
            # If you really wanted to, you could also have subtypes in the subtypes
            {
                "cef": "destinationDnsDomain",
                "name": "Domain Artifact",
                "callback": _extract_domain_from_url,
            }  # Method to extract substring
        ]
    }

    EMAILS_DOMAIN_SUBTYPES_DICT = {
        "subtypes": [
            {
                "cef": "destinationDnsDomain",
                "name": "Domain Artifact",
                "callback": lambda x: x[x.rfind("@") + 1 :],
                "validator": lambda x: not _is_ip(x),
            }
        ]
    }

    found_values = set()

    def __init__(self, parse_domains: bool, patterns: Optional[list[dict[str, Any]]] = None):
        self.patterns = self.BASE_PATTERNS if patterns is None else patterns

        if parse_domains:
            # Add the subtypes somain parsing functions only if parse_domains is True
            is_email = True
            for pattern_dict in self.patterns:
                if pattern_dict.get("cef") == "requestURL" and pattern_dict.get("pattern") == URI_REGEX:
                    pattern_dict.update(self.URL_DOMAIN_SUBTYPES_DICT)
                    is_email = False
                elif pattern_dict.get("cef") == "email" and pattern_dict.get("pattern") in [EMAIL_REGEX, EMAIL_REGEX2]:
                    pattern_dict.update(self.EMAILS_DOMAIN_SUBTYPES_DICT)
            if is_email:
                self.patterns.append(self.DOMAIN_PATTERN)
        self.added_artifacts = 0

    def _create_artifact(self, artifacts: list[Artifact], value: Any, cef: str, name: str) -> None:
        artifact = Artifact(
            source_data_identifier=self.added_artifacts,
            cef={cef: value},
            name=name,
        )
        artifacts.append(artifact)
        self.added_artifacts += 1
        self.found_values.add(value)

    def _parse_ioc_subtype(self, artifacts: list[Artifact], value: Any, subtype: dict[str, Any]) -> None:
        callback = subtype.get("callback")
        if callback:
            sub_val = callback(value)
            self._pass_over_value(artifacts, sub_val, subtype)

    def _pass_over_value(self, artifacts: list[Artifact], value: Any, ioc: dict[str, Any]) -> None:
        validator = ioc.get("validator")
        clean = ioc.get("clean")
        subtypes = ioc.get("subtypes", [])
        if not value:
            return
        if value in self.found_values:
            return
        if clean:
            value = clean(value)
        if validator and not validator(value):
            return
        self._create_artifact(artifacts, value, ioc["cef"], ioc["name"])
        for st in subtypes:
            self._parse_ioc_subtype(artifacts, value, st)

    def parse_to_artifacts(self, text: str) -> list[Artifact]:
        artifacts = []
        for ioc in self.patterns:
            found = re.findall(ioc["pattern"], text, flags=re.IGNORECASE)
            for match in found:
                if isinstance(match, tuple):
                    for x in match:
                        self._pass_over_value(artifacts, x, ioc)
                else:
                    self._pass_over_value(artifacts, match, ioc)
        return artifacts

    def add_artifact(self, text: str) -> Artifact:
        artifact = Artifact(
            source_data_identifier=self.added_artifacts,
            cef={"message": text},
            name="Raw Text Artifact",
        )
        self.added_artifacts += 1
        self.found_values.add(text)
        return artifact


def _grab_raw_text(action_result: "ActionResult", txt_file: str) -> tuple[bool, Optional[str]]:
    """This function will actually really work for any file which is basically raw text.
    html, rtf, and the list could go on
    """
    try:
        with open(txt_file, "rb") as fp:
            text = UnicodeDammit(fp.read()).unicode_markup
        return phantom.APP_SUCCESS, text
    except Exception as e:
        error_code, error_message = _get_error_message_from_exception(e)
        error_text = f"Error Code: {error_code}. Error Message: {error_message}"
        return action_result.set_status(phantom.APP_ERROR, error_text), None


class PDFXrefObjectsToXML:
    """
    Class contains the methods to Convert the PDF cross reference table(xref) objects to XML
    The xref is the index by which all of the indirect objects, in the PDF file are located.
    https://labs.appligent.com/pdfblog/pdf_cross_reference_table/
    """

    @classmethod
    def encode(cls, data: bytes) -> str:
        """Encode characters of text"""
        buf = StringIO()
        for byte in data:
            if byte < 32 or 127 <= byte or byte in ESCAPE:
                buf.write(f"&#{byte};")
            else:
                buf.write(chr(byte))
        return buf.getvalue()

    @classmethod
    def dump_xml(cls, text: str, obj: Any) -> str:
        """Convert PDF xref object to XML"""
        if obj is None:
            text += "<null />"
            return text

        if isinstance(obj, dict):
            text += f'<dict size="{len(obj)}">\n'
            for key, value in obj.items():
                text += f"<key>\n{key}\n</key>\n"
                text += "<value>"
                text = cls.dump_xml(text, value)
                text += "</value>\n"
            text += "</dict>"
            return text

        if isinstance(obj, list):
            text += f'<list size="{len(obj)}">\n'
            for value in obj:
                text = cls.dump_xml(text, value)
                text += "\n"
            text += "</list>"
            return text

        if isinstance(obj, bytes):
            text += f'<string size="{len(obj)}">\n{cls.encode(obj)}\n</string>'
            return text

        if isinstance(obj, PDFStream):
            text += "<stream>\n<props>\n"
            text = cls.dump_xml(text, obj.attrs)
            text += "\n</props>\n"
            text += "</stream>"
            return text

        if isinstance(obj, PDFObjRef):
            text += f'<ref id="{obj.objid}" />'
            return text

        if isinstance(obj, PSKeyword):
            text += f"<keyword>\n{obj.name}\n</keyword>"
            return text

        if isinstance(obj, PSLiteral):
            text += f"<literal>\n{obj.name}\n</literal>"
            return text

        if isnumber(obj):
            text += f"<number>\n{obj}\n</number>"
            return text

        raise TypeError(f"Unable to extract the object from PDF. Reason: {obj}")

    @classmethod
    def dump_trailers(cls, text: str, doc: PDFDocument) -> str:
        """Iterate trough xrefs and convert trailer of xref to XML"""
        for xref in doc.xrefs:
            if trailer := getattr(xref, "trailer"):
                text += "<trailer>\n"
                cls.dump_xml(text, trailer)
                text += "\n</trailer>\n\n"
        return text

    @classmethod
    def convert_objects_to_xml_text(cls, text: str, doc: PDFDocument) -> str:
        """Iterate trough xrefs and convert objects of xref to XML"""
        visited = set()
        text += "<pdf>"
        for xref in doc.xrefs:
            for obj_id in xref.get_objids():
                if obj_id in visited:
                    continue
                visited.add(obj_id)
                try:
                    obj = doc.getobj(obj_id)
                    if obj is None:
                        continue
                    text += f'<object id="{obj_id}">\n'
                    text = cls.dump_xml(text, obj)
                    text += "\n</object>\n\n"
                except PDFObjectNotFound as e:
                    raise PDFObjectNotFound(f"While converting PDF to xml objects PDF object not found. Reason: {e}")
        cls.dump_trailers(text, doc)
        text += "</pdf>"
        return text

    @classmethod
    def pdf_xref_objects_to_xml(cls, pdf_file: str) -> str:
        """Converts PDF cross reference table(xref) objects to XML
        The xref is the index by which all of the indirect objects, in the PDF file are located.
        https://labs.appligent.com/pdfblog/pdf_cross_reference_table/
        """
        text = ""
        with open(pdf_file, "rb") as fp:
            parser = PDFParser(fp)
            doc = PDFDocument(parser)
            text = cls.convert_objects_to_xml_text(text, doc)
        return text


def _pdf_to_text(action_result: "ActionResult", pdf_file: str) -> tuple[bool, Optional[str]]:
    try:
        pagenums = set()
        output = StringIO()
        manager = PDFResourceManager()
        converter = TextConverter(manager, output, laparams=LAParams())
        interpreter = PDFPageInterpreter(manager, converter)
        with open(pdf_file, "rb") as infile:
            for page in PDFPage.get_pages(infile, pagenums):
                interpreter.process_page(page)
        converter.close()
        text = output.getvalue()
        output.close()
        text += PDFXrefObjectsToXML.pdf_xref_objects_to_xml(pdf_file)
        return phantom.APP_SUCCESS, text
    except PDFPasswordIncorrect:
        return (
            action_result.set_status(
                phantom.APP_ERROR,
                "Failed to parse pdf: The provided pdf is password protected",
            ),
            None,
        )
    except PDFEncryptionError:
        return action_result.set_status(phantom.APP_ERROR, "Failed to parse pdf: The provided pdf is encrypted"), None
    except struct.error:
        return (
            action_result.set_status(
                phantom.APP_ERROR,
                "Failed to parse pdf: The provided pdf is password protected or is in different format",
            ),
            None,
        )
    except Exception as e:
        error_code, error_message = _get_error_message_from_exception(e)
        error_text = f"Error Code: {error_code}. Error Message: {error_message}"
        return action_result.set_status(phantom.APP_ERROR, f"Failed to parse pdf: {error_text}"), None


def _docx_to_text(action_result: "ActionResult", docx_file: str) -> tuple[bool, Optional[str]]:
    try:
        doc = docx.Document(docx_file)
    except zipfile.BadZipfile:
        return (
            action_result.set_status(
                phantom.APP_ERROR,
                "Failed to parse docx: The file might be corrupted or password protected or not a docx file",
            ),
            None,
        )
    except Exception as e:
        error_code, error_message = _get_error_message_from_exception(e)
        error_text = f"Error Code: {error_code}. Error Message: {error_message}"
        return action_result.set_status(phantom.APP_ERROR, f"Failed to parse docx: {error_text}"), None

    full_text = []

    # First, render the text in the doc into a string
    for paragraph in doc.paragraphs:
        para_text = "".join(run.text.strip() for run in paragraph.runs).strip()
        # Add the processed paragraph to the full text
        if para_text:
            full_text.append(para_text)

    # Next, expand and append relationship targets present in the document, for searching later
    for rel in doc.part.rels.values():
        # Simple hyperlink, make sure its target is present in the text
        if rel.reltype == REL_TYPE.HYPERLINK:
            full_text.append(rel._target)
        # This is like an embedded HTML within a docx file, stored as bytes
        elif rel.reltype == REL_TYPE.A_F_CHUNK:
            target = cast(DocxPart, rel._target)
            full_text.extend(target.blob.decode(errors="replace").splitlines())

    return phantom.APP_SUCCESS, "\n".join(full_text)


def _csv_to_text(action_result: "ActionResult", csv_file: str) -> tuple[bool, Optional[str]]:
    """This function really only exists due to a misunderstanding on how word boundaries (\b) work
    As it turns out, only word characters can invalidate word boundaries. So stuff like commas,
    brackets, gt and lt signs, etc. do not
    """
    text = ""
    try:
        with open(csv_file) as fp:
            reader = csv.reader(fp)
            for row in reader:
                text += " ".join(row)
                text += " "  # The humanity of always having a trailing space

        return phantom.APP_SUCCESS, text
    except Exception as e:
        error_code, error_message = _get_error_message_from_exception(e)
        error_text = f"Error Code: {error_code}. Error Message: {error_message}"
        return action_result.set_status(phantom.APP_ERROR, f"Failed to parse csv: {error_text}"), None


def _html_to_text(
    action_result: "ActionResult",
    html_file: Optional[str],
    text_val: Optional[str] = None,
) -> tuple[bool, Optional[str]]:
    """Similar to CSV, this is also unnecessary. It will trim /some/ of that fat from a normal HTML, however"""
    try:
        if text_val is None and html_file is not None:
            with open(html_file, "rb") as fp:
                html_text = UnicodeDammit(fp.read()).unicode_markup
        else:
            html_text = text_val

        # To unescape html escaped body
        html_text = unescape(html_text or "")

        soup = BeautifulSoup(html_text, "html.parser")
        read_text = soup.findAll(text=True)
        links = [tag.get("href") for tag in soup.findAll(href=True)]
        srcs = [tag.get("src") for tag in soup.findAll(src=True)]
        text = " ".join(read_text + links + srcs)
        return phantom.APP_SUCCESS, text
    except Exception as e:
        error_code, error_message = _get_error_message_from_exception(e)
        error_text = f"Error Code: {error_code}. Error Message: {error_message}"
        return action_result.set_status(phantom.APP_ERROR, f"Failed to parse html: {error_text}"), None


def _join_thread(base_connector: "BaseConnector", thread: threading.Thread) -> None:
    base_connector._lock.acquire()
    base_connector._done = True
    base_connector._lock.release()
    thread.join()


def _wait_for_parse(base_connector: "BaseConnector") -> None:
    i = 0
    base_msg = "Parsing PDF document"
    while True:
        base_connector._lock.acquire()
        if base_connector._done:
            base_connector._lock.release()
            break
        base_connector.send_progress(base_msg + "." * i)
        base_connector._lock.release()
        i = i % 5 + 1
        time.sleep(1)
    return


def parse_file(
    base_connector: "BaseConnector",
    action_result: "ActionResult",
    file_info: FileInfo,
    parse_domains: bool = True,
    keep_raw: bool = False,
) -> tuple[bool, Optional[dict[str, list[Artifact]]]]:
    """Parse a non-email file"""

    try:
        tiocp = TextIOCParser(parse_domains)
    except Exception as e:
        return action_result.set_status(phantom.APP_ERROR, str(e)), None

    raw_text = None
    if file_info["type"] == "pdf":
        """Parsing a PDF document over like, 10 pages starts to take a while
        (A 80 something page document took like 5 - 10 minutes)
        The thread is nice because it shows a constantly changing message,
        which shows that the app isn't frozen, but it also stops watchdog
        from terminating the app
        """
        thread = threading.Thread(target=_wait_for_parse, args=[base_connector])
        thread.start()
        ret_val, raw_text = _pdf_to_text(action_result, file_info["path"])
        _join_thread(base_connector, thread)
    elif file_info["type"] == "txt":
        ret_val, raw_text = _grab_raw_text(action_result, file_info["path"])
    elif file_info["type"] == "docx":
        ret_val, raw_text = _docx_to_text(action_result, file_info["path"])
    elif file_info["type"] == "csv":
        ret_val, raw_text = _csv_to_text(action_result, file_info["path"])
    elif file_info["type"] == "html":
        ret_val, raw_text = _html_to_text(action_result, file_info["path"])
    else:
        return action_result.set_status(phantom.APP_ERROR, "Unexpected file type"), None

    if phantom.is_fail(ret_val) or raw_text is None:
        return ret_val, None

    base_connector.save_progress("Parsing for IOCs")
    try:
        artifacts = tiocp.parse_to_artifacts(raw_text)
        if keep_raw:
            base_connector.save_progress("Saving Raw Text")
            artifacts.append(tiocp.add_artifact(raw_text))
    except Exception as e:
        error_code, error_message = _get_error_message_from_exception(e)
        error_text = f"Error Code: {error_code}. Error Message: {error_message}"
        return action_result.set_status(phantom.APP_ERROR, error_text), None
    return phantom.APP_SUCCESS, {"artifacts": artifacts}


def parse_structured_file(action_result: "ActionResult", file_info: FileInfo) -> tuple[bool, Optional[dict[str, list[Artifact]]]]:
    if file_info["type"] == "csv":
        csv_file = file_info["path"]
        artifacts = []
        try:
            with open(csv_file) as fp:
                reader = csv.DictReader(fp, restkey="other")  # need to handle lines terminated in commas
                for row in reader:
                    row["source_file"] = file_info["name"]
                    artifacts.append(
                        {
                            "name": "CSV entry",
                            "cef": {k: v for k, v in list(row.items())},
                        }
                    )  # make CSV entry artifact
        except Exception as e:
            error_code, error_message = _get_error_message_from_exception(e)
            error_text = f"Error Code: {error_code}. Error Message: {error_message}"
            return (
                action_result.set_status(
                    phantom.APP_ERROR,
                    f"Failed to parse structured CSV: {error_text}",
                ),
                None,
            )
    else:
        return action_result.set_status(phantom.APP_ERROR, "Structured extraction only supported for CSV files"), None
    return phantom.APP_SUCCESS, {"artifacts": artifacts}


def parse_text(
    base_connector: "BaseConnector",
    action_result: "ActionResult",
    file_type: Optional[str],
    text_val: str,
    parse_domains: bool = True,
) -> tuple[bool, Optional[dict[str, list[Artifact]]]]:
    """Parse a non-email file"""

    try:
        tiocp = TextIOCParser(parse_domains)
    except Exception as e:
        return action_result.set_status(phantom.APP_ERROR, str(e)), None

    raw_text = None
    if file_type == "html":
        ret_val, raw_text = _html_to_text(action_result, None, text_val=text_val)
    elif file_type == "txt" or file_type == "csv":
        ret_val, raw_text = phantom.APP_SUCCESS, text_val
    else:
        return action_result.set_status(phantom.APP_ERROR, "Unexpected file type"), None

    if phantom.is_fail(ret_val) or raw_text is None:
        return ret_val, None

    base_connector.save_progress("Parsing for IOCs")
    try:
        artifacts = tiocp.parse_to_artifacts(raw_text)
    except Exception as e:
        error_code, error_message = _get_error_message_from_exception(e)
        error_text = f"Error Code: {error_code}. Error Message: {error_message}"
        return action_result.set_status(phantom.APP_ERROR, error_text), None

    return phantom.APP_SUCCESS, {"artifacts": artifacts}
